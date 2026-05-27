"""
Certificate document generation module using python-docx
Generates certificates from Word templates with placeholder replacement and signature support.
"""

import os
from datetime import datetime
from io import BytesIO

from bs4 import BeautifulSoup
from django.conf import settings
from django.core.files.base import ContentFile
from django.template import Context, Template
from docx import Document

from .models import GeneratedCertificate


class CertificateDocumentGenerator:
    """
    Generate certificate Word documents from templates with dynamic content and optional signatures.
    Uses python-docx to manipulate .docx templates.
    """
    
    def __init__(self, certificate: GeneratedCertificate):
        self.certificate = certificate
        self.template = certificate.template
        self.data = certificate.certificate_data.copy()
        self.signature = certificate.signature_official if certificate.include_signature else None
    
    def _prepare_data(self):
        """Add computed fields and missing data before replacement"""
        # Add recipient_name from certificate model
        self.data['recipient_name'] = self.certificate.recipient_name
        
        # Add current date fields
        now = datetime.now()
        self.data['day'] = now.strftime('%d')
        self.data['month'] = now.strftime('%B')
        self.data['year'] = now.strftime('%Y')
    
    def generate_from_template(self, template_path):
        """
        Generate certificate by loading template and replacing placeholders.
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        # Prepare data
        self._prepare_data()
        
        # Open template
        doc = Document(template_path)
        
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            # Get combined text from all runs
            combined_text = ''.join(run.text for run in paragraph.runs)
            new_text = combined_text
            
            # Replace all placeholders
            for key, value in self.data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, str(value) if value else '')
            
            # Update if changed
            if new_text != combined_text:
                # Clear all runs
                for run in paragraph.runs:
                    run.text = ''
                # Add new text
                paragraph.add_run(new_text)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        combined_text = ''.join(run.text for run in paragraph.runs)
                        new_text = combined_text
                        for key, value in self.data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in new_text:
                                new_text = new_text.replace(placeholder, str(value) if value else '')
                        if new_text != combined_text:
                            for run in paragraph.runs:
                                run.text = ''
                            paragraph.add_run(new_text)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_from_html_template(self):
        """
        Generate a simple Word document from the HTML template when the original
        .docx template file is unavailable.
        """
        self._prepare_data()

        template = Template(self.template.html_template or "")
        rendered_html = template.render(Context(self.data))
        soup = BeautifulSoup(rendered_html, "html.parser")

        doc = Document()
        doc.add_heading(self.template.template_name, level=1)
        doc.add_paragraph(f"Recipient: {self.certificate.recipient_name}")
        doc.add_paragraph("")

        content_nodes = soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "br", "table"])
        if not content_nodes:
            plain_text = soup.get_text("\n", strip=True)
            if plain_text:
                for line in plain_text.splitlines():
                    if line.strip():
                        doc.add_paragraph(line.strip())
        else:
            for node in content_nodes:
                if node.name in {"h1", "h2", "h3", "h4"}:
                    level = min(int(node.name[1]), 4)
                    text = node.get_text(" ", strip=True)
                    if text:
                        doc.add_heading(text, level=level)
                elif node.name == "p":
                    text = node.get_text(" ", strip=True)
                    if text:
                        doc.add_paragraph(text)
                elif node.name == "li":
                    text = node.get_text(" ", strip=True)
                    if text:
                        doc.add_paragraph(text, style="List Bullet")
                elif node.name == "br":
                    doc.add_paragraph("")
                elif node.name == "table":
                    rows = node.find_all("tr")
                    if not rows:
                        continue
                    col_count = max(len(row.find_all(["td", "th"])) for row in rows)
                    if col_count == 0:
                        continue
                    table = doc.add_table(rows=len(rows), cols=col_count)
                    for row_index, row in enumerate(rows):
                        cells = row.find_all(["td", "th"])
                        for col_index, cell in enumerate(cells):
                            table.cell(row_index, col_index).text = cell.get_text(" ", strip=True)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def save_generated_document(self, template_path, certificate_obj=None):
        """Generate document and save to certificate."""
        if certificate_obj is None:
            certificate_obj = self.certificate
        
        try:
            doc_buffer = self.generate_from_template(template_path)
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            recipient_name = self.certificate.recipient_name.replace(' ', '_')[:30]
            template_name = self.template.template_type
            docx_filename = f"{template_name}_{recipient_name}_{timestamp}.docx"
            
            certificate_obj.docx_file.save(docx_filename, ContentFile(doc_buffer.getvalue()), save=True)
            return certificate_obj.docx_file.path
            
        except Exception as e:
            raise e
    
    def save_files(self, certificate_obj=None):
        """Save certificate using template file."""
        if certificate_obj is None:
            certificate_obj = self.certificate

        template_path = None
        if self.template.template_file:
            try:
                template_path = self.template.template_file.path
            except (ValueError, NotImplementedError):
                template_path = None

        if template_path and os.path.exists(template_path):
            return self.save_generated_document(template_path, certificate_obj)

        doc_buffer = self.generate_from_html_template()
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        recipient_name = self.certificate.recipient_name.replace(' ', '_')[:30]
        template_name = self.template.template_type
        docx_filename = f"{template_name}_{recipient_name}_{timestamp}.docx"
        certificate_obj.docx_file.save(docx_filename, ContentFile(doc_buffer.getvalue()), save=True)
        return certificate_obj.docx_file.path
