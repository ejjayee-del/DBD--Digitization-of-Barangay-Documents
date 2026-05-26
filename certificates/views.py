# Standard library imports
from datetime import datetime
from io import BytesIO

# Third-party imports
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

# Django imports
from django import forms
from django.contrib import messages
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.http import FileResponse, HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.db.models import Q
from django.utils import timezone
from django.views.decorators.http import require_http_methods
from django_summernote.widgets import SummernoteWidget

# Local imports
from transactions.models import ActivityLog
from .models import CertificateDocument, CertificateTemplate, GeneratedCertificate, CertificateRequest
from .forms import CertificateGenerationForm, DynamicCertificateForm


class EditableCertificateForm(forms.ModelForm):
    """Form for editing certificate content in the system"""
    class Meta:
        model = GeneratedCertificate
        fields = ['certificate_data']
        widgets = {
            'certificate_data': SummernoteWidget(),
        }


def _prepare_requester_dynamic_form(dynamic_form):
    """Remove requester-only dynamic fields and auto-fill the date field."""
    for field_name in ('day', 'month', 'year', 'number', 'amount_paid', 'date_paid'):
        if field_name in dynamic_form.fields:
            del dynamic_form.fields[field_name]

    if 'date' in dynamic_form.fields:
        date_field = dynamic_form.fields['date']
        dynamic_form.fields['date'] = forms.DateField(
            label=date_field.label,
            required=date_field.required,
            help_text=date_field.help_text,
            initial=timezone.now().date(),
            widget=forms.DateInput(attrs={
                'class': 'form-control',
                'type': 'date',
                'readonly': 'readonly'
            })
        )


@login_required
@require_http_methods(["GET"])
def certificate_types_view(request):
    """Display all available certificate types"""
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to generate certificates.')
        return redirect('dashboard')
    
    query = request.GET.get('q', '').strip()
    certificates = CertificateTemplate.objects.filter(is_active=True)

    if query:
        certificates = certificates.filter(
            Q(template_name__icontains=query) |
            Q(description__icontains=query) |
            Q(template_type__icontains=query)
        )

    certificates = list(certificates)
    kp_form_types = ['hearing_notice', 'complaint_form', 'settlement']
    
    certificates.sort(key=lambda cert: (0, kp_form_types.index(cert.template_type)) if cert.template_type in kp_form_types else (1, cert.template_name))
    
    context = {
        'certificates': certificates,
        'kp_form_types': kp_form_types,
        'query': query,
    }
    
    return render(request, 'certificates/certificate_types.html', context)


@login_required
@require_http_methods(["GET", "POST"])
def generate_certificate_view(request, template_id):
    """Generate a new certificate from template"""
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to generate certificates.')
        return redirect('dashboard')
    
    template = get_object_or_404(CertificateTemplate, id=template_id, is_active=True)
    
    if request.method == 'POST':
        form = CertificateGenerationForm(request.POST)
        dynamic_form = DynamicCertificateForm(template, request.POST)
        
        if form.is_valid() and dynamic_form.is_valid():
            # Create certificate
            certificate = form.save(commit=False)
            certificate.template = template
            certificate.created_by = request.user
            certificate.status = 'generated'
            
            # Convert Decimal values to strings or floats before storing as JSON
            from decimal import Decimal
            cleaned_data = dynamic_form.cleaned_data
            for key, value in cleaned_data.items():
                if isinstance(value, Decimal):
                    # Convert Decimal to string (preserves precision) OR float (loses precision)
                    cleaned_data[key] = str(value)  # Using string for precision
                    # Alternatively use: cleaned_data[key] = float(value)
            
            certificate.certificate_data = cleaned_data
            certificate.save()
            
            # Generate PDF and Word files - lazy import
            try:
                from .document_generator import CertificateDocumentGenerator
                generator = CertificateDocumentGenerator(certificate)
                generator.save_files(certificate)
            except ImportError as e:
                # Log the exception and notify user
                import logging
                logger = logging.getLogger(__name__)
                logger.error(f"Document generation failed: {e}")
                messages.warning(request, 'Certificate saved but document generation is currently unavailable. Please contact an administrator.')
            
            # Log the activity
            ActivityLog.objects.create(
                user=request.user,
                log_type='certificate_generated',
                action_description=f"Generated {template.template_name} for {certificate.recipient_name}",
                object_id=certificate.id,
            )
            
            messages.success(request, 'Certificate generated successfully!')
            return redirect('certificates:preview', certificate_id=certificate.id)
    else:
        form = CertificateGenerationForm()
        dynamic_form = DynamicCertificateForm(template)
    
    context = {
        'template': template,
        'form': form,
        'dynamic_form': dynamic_form,
    }
    
    return render(request, 'certificates/generate_certificate.html', context)


@login_required
@require_http_methods(["GET"])
def preview_certificate_view(request, certificate_id):
    """Preview generated certificate"""
    certificate = get_object_or_404(GeneratedCertificate, id=certificate_id)
    
    # Check permission
    is_staff = request.user.can_view_history
    is_requester_owner = (request.user.role == 'requester' and 
                          certificate.certificate_request and 
                          certificate.certificate_request.requester == request.user)
    is_creator = certificate.created_by == request.user
    
    if not (is_staff or is_requester_owner or is_creator):
        messages.error(request, 'You do not have permission to view this certificate.')
        return redirect('dashboard')
    
    # Render the HTML template with certificate data
    from django.template import Template, Context
    template = Template(certificate.template.html_template)
    context_data = Context(certificate.certificate_data)
    rendered_content = template.render(context_data)
    
    context = {
        'certificate': certificate,
        'rendered_content': rendered_content,
    }
    
    # Use different templates for staff vs requester
    if request.user.role == 'requester':
        # Requesters see view-only template (no edit/print/release buttons)
        return render(request, 'certificates/view_certificate_only.html', context)
    else:
        # Staff/Admin see full preview with all action buttons
        return render(request, 'certificates/preview_certificate.html', context)


@login_required
@require_http_methods(["GET"])
def download_certificate_view(request, certificate_id, file_type='docx'):
    """Download certificate as Word document"""
    certificate = get_object_or_404(GeneratedCertificate, id=certificate_id)
    
    # Check permission
    if certificate.created_by != request.user and not request.user.can_view_history:
        messages.error(request, 'You do not have permission to view this certificate.')
        return redirect('dashboard')
    
    if certificate.docx_file:
        file_path = certificate.docx_file.path
        f = open(file_path, 'rb')
        return FileResponse(f, 
                          content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                          as_attachment=True,
                          filename=f"{certificate.recipient_name}_{certificate.template.template_name}.docx")
    
    messages.error(request, 'Certificate file not found.')
    return redirect('certificates:preview', certificate_id=certificate.id)

@login_required
@require_http_methods(["GET"])
def print_certificate_view(request, certificate_id):
    """Mark certificate as printed and record it"""
    certificate = get_object_or_404(GeneratedCertificate, id=certificate_id)
    
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to print certificates.')
        return redirect('dashboard')
    
    # Update certificate status
    certificate.status = 'printed'
    certificate.printed_date = timezone.now()
    certificate.printed_by = request.user
    certificate.save()
    
    # Log the activity
    ActivityLog.objects.create(
        user=request.user,
        log_type='certificate_printed',
        action_description=f"Printed {certificate.template.template_name} for {certificate.recipient_name}",
        object_id=certificate.id,
    )
    
    messages.success(request, 'Certificate marked as printed.')
    return redirect('certificates:preview', certificate_id=certificate_id)


@login_required
@require_http_methods(["GET"])
def list_certificates_view(request):
    """List all generated certificates"""
    if not request.user.can_view_history:
        messages.error(request, 'You do not have permission to view certificates.')
        return redirect('dashboard')
    
    certificates = GeneratedCertificate.objects.select_related('template', 'created_by').all()
    
    # Filter by status if provided
    status = request.GET.get('status')
    if status:
        certificates = certificates.filter(status=status)
    
    # Pagination
    paginator = Paginator(certificates, 25)  # 25 per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'certificates': page_obj,
    }
    
    return render(request, 'certificates/list_certificates.html', context)


@staff_member_required
def download_document_as_docx(request, document_id):
    """Download certificate document as Word file"""
    document = get_object_or_404(CertificateDocument, id=document_id)
    
    # Create Word document
    doc = Document()
    
    # Add title
    doc.add_heading(document.title, 0)
    
    # Parse HTML content from Summernote
    soup = BeautifulSoup(document.content, 'html.parser')
    
    # Convert HTML to Word document elements
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'table', 'ul', 'ol']):
        if element.name == 'p':
            p = doc.add_paragraph()
            p.add_run(element.get_text())
        elif element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'table':
            # Convert HTML table to Word table
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text()
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{document.title}_{document.id}.docx"'
    return response


@staff_member_required
def certificate_list(request):
    """List all certificates for the Generate page"""
    documents = CertificateDocument.objects.all().order_by('-updated_at')
    return render(request, 'certificates/certificate_list.html', {'documents': documents})


@login_required
@require_http_methods(["GET", "POST"])
def edit_certificate_view(request, certificate_id):
    """Edit certificate content in the system before printing"""
    certificate = get_object_or_404(GeneratedCertificate, id=certificate_id)
    
    # Check permission
    if certificate.created_by != request.user and not request.user.can_edit_records:
        messages.error(request, 'You do not have permission to edit this certificate.')
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = EditableCertificateForm(request.POST, instance=certificate)
        if form.is_valid():
            form.save()
            
            # Re-generate the Word document with edited content
            try:
                from .document_generator import CertificateDocumentGenerator
                generator = CertificateDocumentGenerator(certificate)
                generator.save_files(certificate)
            except ImportError:
                pass
            
            messages.success(request, 'Certificate updated successfully!')
            return redirect('certificates:preview', certificate_id=certificate.id)
    else:
        form = EditableCertificateForm(instance=certificate)
    
    # Convert certificate_data to HTML for editing if it's a dict
    if isinstance(certificate.certificate_data, dict):
        # Format the data as readable text
        html_content = '<div class="certificate-content">'
        for key, value in certificate.certificate_data.items():
            if value:
                html_content += f'<p><strong>{key.replace("_", " ").title()}:</strong> {value}</p>'
        html_content += '</div>'
        form.initial['certificate_data'] = html_content
    
    context = {
        'certificate': certificate,
        'form': form,
    }
    
    return render(request, 'certificates/edit_certificate.html', context)



@login_required
@require_http_methods(["GET"])
def print_certificate_from_system(request, certificate_id):
    """Display certificate as HTML for printing from the system"""
    certificate = get_object_or_404(GeneratedCertificate, id=certificate_id)
    
    # Check permission
    if certificate.created_by != request.user and not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to print this certificate.')
        return redirect('dashboard')
    
    # Render the HTML template with certificate data
    from django.template import Template, Context
    template = Template(certificate.template.html_template)
    context = Context(certificate.certificate_data)
    rendered_content = template.render(context)
    
    context = {
        'certificate': certificate,
        'data': certificate.certificate_data,
        'recipient_name': certificate.recipient_name,
        'certificate_id': certificate.id,
        'rendered_content': rendered_content,
    }
    
    return render(request, 'certificates/print_certificate.html', context)


# Requester views
@login_required
@require_http_methods(["GET"])
def request_certificate_types_view(request):
    """Display available certificate types for requesters"""
    if request.user.role != 'requester':
        messages.error(request, 'This page is for requesters only.')
        return redirect('dashboard')
    
    certificates = CertificateTemplate.objects.filter(is_active=True)
    
    context = {
        'certificates': certificates,
    }
    
    return render(request, 'certificates/request_types.html', context)


@login_required
@require_http_methods(["GET", "POST"])
def request_certificate_view(request, template_id):
    """Request a new certificate"""
    if request.user.role != 'requester':
        messages.error(request, 'This page is for requesters only.')
        return redirect('dashboard')
    
    template = get_object_or_404(CertificateTemplate, id=template_id, is_active=True)
    
    if request.method == 'POST':
        form = CertificateGenerationForm(request.POST)
        dynamic_form = DynamicCertificateForm(template, request.POST)
        _prepare_requester_dynamic_form(dynamic_form)

        if request.POST.get('confirm') and form.is_valid() and dynamic_form.is_valid():
            request_data = dynamic_form.cleaned_data.copy()
            
            # Convert Decimal values to string for JSON serialization
            from decimal import Decimal
            for key, value in request_data.items():
                if isinstance(value, Decimal):
                    request_data[key] = str(value)
            
            now = timezone.now()
            request_data['date'] = now.strftime('%B %d, %Y')
            request_data['day'] = now.strftime('%d')
            request_data['month'] = now.strftime('%B')
            request_data['year'] = now.strftime('%Y')
            request_data.setdefault('number', '')
            request_data.setdefault('amount_paid', '')
            request_data.setdefault('date_paid', '')

            cert_request = CertificateRequest.objects.create(
                template=template,
                requester=request.user,
                recipient_name=form.cleaned_data['recipient_name'],
                recipient_email=form.cleaned_data.get('recipient_email', ''),
                recipient_contact=form.cleaned_data.get('recipient_contact', ''),
                request_data=request_data,
                status='pending'
            )

            from transactions.models import VisitHistory
            VisitHistory.objects.create(
                requester=request.user,
                certificate_request=cert_request,
                visit_type='online_request',
                purpose_of_visit=f"Certificate Request: {template.template_name}",
            )
            
            ActivityLog.objects.create(
                user=request.user,
                log_type='certificate_requested',
                action_description=f"Requested {template.template_name} for {cert_request.recipient_name}",
                object_id=cert_request.id,
            )
            
            messages.success(request, 'Certificate request submitted successfully!')
            return redirect('certificates:my_requests')

        if request.POST.get('edit'):
            # Re-bind forms so the user can correct mistakes
            form = CertificateGenerationForm(request.POST)
            dynamic_form = DynamicCertificateForm(template, request.POST)
            _prepare_requester_dynamic_form(dynamic_form)

        elif form.is_valid() and dynamic_form.is_valid():
            preview_data = dynamic_form.cleaned_data.copy()
            
            # Convert Decimal values to string for JSON serialization
            from decimal import Decimal
            for key, value in preview_data.items():
                if isinstance(value, Decimal):
                    preview_data[key] = str(value)
            
            now = timezone.now()
            preview_data.setdefault('date', now.strftime('%B %d, %Y'))
            preview_data.setdefault('day', now.strftime('%d'))
            preview_data.setdefault('month', now.strftime('%B'))
            preview_data.setdefault('year', now.strftime('%Y'))
            preview_data.setdefault('number', '')
            preview_data.setdefault('amount_paid', '')
            preview_data.setdefault('date_paid', '')

            from django.template import Template, Context
            render_ctx = Context({**preview_data, 'recipient_name': form.cleaned_data.get('recipient_name')})
            template_obj = Template(template.html_template)
            rendered_content = template_obj.render(render_ctx)

            context = {
                'template': template,
                'form': form,
                'dynamic_form': dynamic_form,
                'rendered_content': rendered_content,
                'posted_data': request.POST,
            }
            return render(request, 'certificates/preview_request.html', context)
    else:
        form = CertificateGenerationForm()
        dynamic_form = DynamicCertificateForm(template)
        _prepare_requester_dynamic_form(dynamic_form)
    
    context = {
        'template': template,
        'form': form,
        'dynamic_form': dynamic_form,
    }
    
    return render(request, 'certificates/request_certificate.html', context)


@login_required
@require_http_methods(["GET"])
def my_requests_view(request):
    """List user's certificate requests"""
    if request.user.role != 'requester':
        messages.error(request, 'This page is for requesters only.')
        return redirect('dashboard')

    requests = CertificateRequest.objects.filter(requester=request.user).order_by('-created_date')

    context = {
        'requests': requests,
    }

    return render(request, 'certificates/my_requests.html', context)

@login_required
@require_http_methods(["GET", "POST"])
def release_certificate_view(request, certificate_id):
    """Release a certificate to the requester (mark as released/completed)"""
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to release certificates.')
        return redirect('dashboard')
    
    certificate = get_object_or_404(GeneratedCertificate, id=certificate_id)
    
    if request.method == 'POST':
        # Mark certificate as released
        certificate.status = 'released'
        certificate.released_date = timezone.now()
        certificate.released_by = request.user
        certificate.save()
        
        # Update related certificate request if exists
        try:
            cert_request = CertificateRequest.objects.get(generated_certificate=certificate)
            cert_request.status = 'completed'
            cert_request.save()
        except CertificateRequest.DoesNotExist:
            pass
        
        # Log the activity
        ActivityLog.objects.create(
            user=request.user,
            log_type='certificate_released',
            action_description=f"Released {certificate.template.template_name} for {certificate.recipient_name}",
            object_id=certificate.id,
        )
        
        messages.success(request, 'Certificate released successfully!')
        return redirect('certificates:manage_requests')
    
    context = {
        'certificate': certificate,
    }
    
    return render(request, 'certificates/release_certificate.html', context)

@login_required
@require_http_methods(["GET"])
def manage_requests_view(request):
    """View and manage certificate requests (staff only)"""
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to manage requests.')
        return redirect('dashboard')
    
    requests = CertificateRequest.objects.select_related('template', 'requester').order_by('-created_date')
    
    # Filter by status if provided
    status = request.GET.get('status')
    if status:
        requests = requests.filter(status=status)
    
    context = {
        'requests': requests,
    }
    
    return render(request, 'certificates/manage_requests.html', context)


@login_required
@require_http_methods(["GET"])
def preview_request_staff_view(request, request_id):
    """Preview a certificate request for staff before approval"""
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to view this request.')
        return redirect('dashboard')
    
    cert_request = get_object_or_404(CertificateRequest, id=request_id)
    
    # Render the HTML template with request data
    from django.template import Template, Context
    template = Template(cert_request.template.html_template)
    context = Context(cert_request.request_data)
    rendered_content = template.render(context)
    
    context = {
        'cert_request': cert_request,
        'rendered_content': rendered_content,
    }
    
    return render(request, 'certificates/preview_request_staff.html', context)


@login_required
@require_http_methods(["POST"])
def approve_request_view(request, request_id):
    """Approve a certificate request and generate certificate"""
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to approve requests.')
        return redirect('dashboard')
    
    cert_request = get_object_or_404(CertificateRequest, id=request_id)
    
    if cert_request.status != 'pending':
        messages.error(request, 'Request has already been processed.')
        return redirect('certificates:manage_requests')
    
    # Update request status
    cert_request.status = 'approved'
    cert_request.reviewed_by = request.user
    cert_request.reviewed_date = timezone.now()
    cert_request.save()
    
    # Generate certificate
    certificate = GeneratedCertificate.objects.create(
        template=cert_request.template,
        recipient_name=cert_request.recipient_name,
        recipient_email=cert_request.recipient_email,
        recipient_contact=cert_request.recipient_contact,
        certificate_data=cert_request.request_data,
        created_by=request.user,
        status='generated'
    )
    
    # Link to request
    cert_request.generated_certificate = certificate
    cert_request.status = 'completed'
    cert_request.save()
    
    # Generate files
    try:
        from .document_generator import CertificateDocumentGenerator
        generator = CertificateDocumentGenerator(certificate)
        generator.save_files(certificate)
    except ImportError:
        pass
    
    # Log activities
    ActivityLog.objects.create(
        user=request.user,
        log_type='certificate_generated',
        action_description=f"Generated certificate from request: {certificate.template.template_name} for {certificate.recipient_name}",
        object_id=certificate.id,
    )
    
    messages.success(request, 'Certificate request approved and certificate generated!')
    return redirect('certificates:manage_requests')


@login_required
@require_http_methods(["POST"])
def reject_request_view(request, request_id):
    """Reject a certificate request"""
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to reject requests.')
        return redirect('dashboard')
    
    cert_request = get_object_or_404(CertificateRequest, id=request_id)
    
    if cert_request.status != 'pending':
        messages.error(request, 'Request has already been processed.')
        return redirect('certificates:manage_requests')
    
    reason = request.POST.get('rejection_reason', '').strip()
    
    # Update request status
    cert_request.status = 'rejected'
    cert_request.reviewed_by = request.user
    cert_request.reviewed_date = timezone.now()
    cert_request.rejection_reason = reason
    cert_request.save()
    
    messages.success(request, 'Certificate request rejected.')
    return redirect('certificates:manage_requests')